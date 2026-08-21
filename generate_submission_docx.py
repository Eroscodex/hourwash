import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = docx.Document()

# Set standard 1-inch margins
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Title
title = doc.add_heading('HOUR WASH SYSTEM: OVERALL ROLE DASHBOARD & SCM QUANTITATIVE ANALYSIS', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('System: Hour Wash Laundry Management System (hourwashweb)\nDashboards Covered: Admin, Staff, Rider, and Customer Portals\n')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Section 1: Dashboard Role Overview
doc.add_heading('SECTION 1: OVERALL SYSTEM DASHBOARD BREAKDOWN (ADMIN, STAFF, RIDER, CUSTOMER)', level=1)

# Admin
doc.add_heading('1. Admin & Manager Dashboard (/admin/dashboard & /admin/analytics)', level=2)
doc.add_paragraph('• Dashboard Role: Strategic Management, Financial Oversight & Store Capacity Control.')
doc.add_paragraph('• Quantitative Analytics Featured:')
doc.add_paragraph('   - Descriptive: Total Revenue = ∑(total_amount | paid), Fleet Utilization % = (Active / Total) × 100%, Total User Count, and Order Pipeline Breakdown (pending, washing, drying, finish, completed).')
doc.add_paragraph('   - Predictive: Store turnaround trends & daily booking volume forecasting.')
doc.add_paragraph('   - Prescriptive: 4:30 PM Cut-off Policy enforcement, store open/closed toggles, and machine maintenance status control.')
doc.add_paragraph('• SCM Functional Area & Process: Property, Facility & Finance Operations — Managing financial records, machine asset health, and staff/rider account allocations.')

# Staff
doc.add_heading('2. Staff & Store Operator Dashboard (/admin/laundry & /admin/machines)', level=2)
doc.add_paragraph('• Dashboard Role: Inbound Processing, Machine Dispatch & Quality Control.')
doc.add_paragraph('• Quantitative Analytics Featured:')
doc.add_paragraph('   - Descriptive: Real-time machine digital timers, live wash cycle monitoring, order queue status per machine.')
doc.add_paragraph('   - Predictive: Live estimated completion countdowns (estimated_minutes) per active order load.')
doc.add_paragraph('   - Prescriptive: Automated machine assignment logic (allocates lowest-workload idle machine) and 12-stamp discount token redemption during walk-in entry.')
doc.add_paragraph('• SCM Functional Area & Process: Facility Operations & Production Intake — Weighing incoming loads (7kg limit), scanning QR code tags, transitioning states (washing -> rinsing -> drying -> finish), and packaging completed orders.')

# Rider
doc.add_heading('3. Rider Logistics Dashboard (/rider/dashboard)', level=2)
doc.add_paragraph('• Dashboard Role: Inbound & Outbound Delivery Dispatch & Doorstep Fulfillment.')
doc.add_paragraph('• Quantitative Analytics Featured:')
doc.add_paragraph('   - Descriptive: Active Pickup Requests Count (N_pickup), Received Count (N_received), Out-for-Delivery Count (N_delivery), and Completed Today Counter.')
doc.add_paragraph('   - Predictive: Arrival time windows based on customer doorstep address and barangay location.')
doc.add_paragraph('   - Prescriptive: Priority dispatch queueing and automated SMS customer notification triggers upon pickup/delivery updates.')
doc.add_paragraph('• SCM Functional Area & Process: Outbound Logistics & Distribution — Collecting laundry from customer locations (out_for_pickup -> received) and delivering finished laundry (out_for_delivery -> completed).')

# Customer
doc.add_heading('4. Customer Portal & Tracking Dashboard (/my-orders & /laundry/track/{order})', level=2)
doc.add_paragraph('• Dashboard Role: Demand Generation, Self-Service Booking & Order Tracking.')
doc.add_paragraph('• Quantitative Analytics Featured:')
doc.add_paragraph('   - Descriptive: Frequent User Stamp Card (X/12), completed cards count, available ₱50.00 discount tokens.')
doc.add_paragraph('   - Predictive: Dynamic completion countdown tracker (estimated_completion timer via QR tag).')
doc.add_paragraph('   - Prescriptive: Automated 12-stamp reward token application (₱50.00 OFF) and Tipid supply choice selection (-₱15.00 / -₱25.00).')
doc.add_paragraph('• SCM Functional Area & Process: Customer Relationship Management (CRM) — Submitting online bookings, tracking live status, redeeming rewards, and rating services.')

# Section 2: Overall Role Matrix Table
doc.add_heading('SECTION 2: OVERALL ROLE DASHBOARD & SCM QUANTITATIVE MATRIX', level=1)

table = doc.add_table(rows=1, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_cells = table.rows[0].cells
hdr_titles = ['Role & Dashboard', 'SCM Functional Area', 'Quantitative Method Featured', 'Key Dashboard Processes']
for i, title in enumerate(hdr_titles):
    hdr_cells[i].text = title
    hdr_cells[i].paragraphs[0].runs[0].font.bold = True

rows_data = [
    ('Admin Dashboard\n(/admin/dashboard)', 'Property, Facility & Finance Operations', 'Descriptive: Total Revenue, Utilization %\nPrescriptive: 4:30 PM Cut-off, Machine Maintenance', '• Monitors revenue & fleet usage\n• Manages machine asset statuses\n• Oversees staff & system parameters'),
    ('Staff Dashboard\n(/admin/laundry)', 'Operations & Facility Intake', 'Descriptive: Active timers, Queue status\nPrescriptive: Auto machine dispatch', '• Weighs laundry (7kg limit)\n• Scans QR tags & controls wash cycles\n• Triggers SMS/Email packaging alerts'),
    ('Rider Dashboard\n(/rider/dashboard)', 'Inbound & Outbound Logistics Distribution', 'Descriptive: Pickup/Delivery task metrics\nPrescriptive: Auto SMS delivery alerts', '• Accepts pickup requests (out_for_pickup)\n• Delivers clean laundry (out_for_delivery)\n• Updates order status to completed'),
    ('Customer Portal\n(/my-orders & /track)', 'CRM & Demand Booking Intake', 'Predictive: Live QR completion ETA\nPrescriptive: 12-Stamp ₱50 OFF token logic', '• Books laundry orders online\n• Scans QR code for live tracking\n• Collects stamps & redeems ₱50 tokens')
]

for role, area, quant, proc in rows_data:
    row_cells = table.add_row().cells
    row_cells[0].text = role
    row_cells[1].text = area
    row_cells[2].text = quant
    row_cells[3].text = proc

output_path = 'Hour_Wash_Quantitative_and_SCM_Assignment_Submission.docx'
doc.save(output_path)
print(f'Successfully generated updated {output_path}')
