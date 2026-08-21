import docx
from docx.shared import Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def create_quant_doc():
    doc = docx.Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    title = doc.add_heading('SUBJECT 1: QUANTITATIVE METHODS IN MANAGEMENT', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('System: Hour Wash Laundry Management System (hourwashweb)\nFocus: Descriptive, Predictive, and Prescriptive Analytics across Dashboard Roles\n').alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('QUANTITATIVE METHODS MATRIX BY DASHBOARD ROLE', level=1)
    
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    hdr[0].text = 'Dashboard Role'
    hdr[1].text = 'Analytics Category'
    hdr[2].text = 'Mathematical Formulas & Models'
    hdr[3].text = 'Quantitative System Logic'
    for c in hdr:
        c.paragraphs[0].runs[0].font.bold = True

    rows = [
        ('Admin Dashboard\n(/admin/dashboard)', 'Descriptive &\nPrescriptive', '• Total Revenue = ∑(total_amount | paid)\n• Fleet Utilization % = (Active / Total) × 100%\n• Cut-off Rule: IF T > 16:30 ⟹ Next Day Queue', 'Calculates real-time financial totals, fleet usage rate %, and enforces 4:30 PM store cut-off rules.'),
        ('Staff Dashboard\n(/admin/laundry)', 'Descriptive, Predictive\n& Prescriptive', '• Machine Timer = remaining_minutes\n• Order Load Count = ceil(Weight / 7.0)\n• Auto Machine Dispatch = min(workload_idle)', 'Monitors live machine countdown timers, calculates load counts based on 7kg limit, and auto-assigns idle machines.'),
        ('Rider Dashboard\n(/rider/dashboard)', 'Descriptive &\nPrescriptive', '• Task Counters: N_pickup, N_received, N_delivery\n• SMS Trigger = Event(Status_Change)', 'Tracks active dispatch task counters and triggers automated SMS notifications upon collection & delivery.'),
        ('Customer Portal\n(/my-orders & /track)', 'Predictive &\nPrescriptive', '• Dynamic ETA = T_creation + t_service + (N_queue * t_cycle)\n• Stamps_new = stamps + 1\n• IF Stamps ≥ 12 ⟹ Rewards + 1 (₱50 OFF), Stamps = 0', 'Forecasts order completion time via dynamic QR countdown and automates 12-stamp card reward resets.')
    ]

    for r1, r2, r3, r4 in rows:
        row = table.add_row().cells
        row[0].text = r1
        row[1].text = r2
        row[2].text = r3
        row[3].text = r4

    doc.save('Hour_Wash_Subject1_Quantitative_Methods.docx')
    print('Generated Hour_Wash_Subject1_Quantitative_Methods.docx')

def create_scm_doc():
    doc = docx.Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    title = doc.add_heading('SUBJECT 2: SUPPLY CHAIN MANAGEMENT (SCM)', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('System: Hour Wash Laundry Management System (hourwashweb)\nFocus: SCM Functional Areas, Business Functions, and Business Processes across Roles\n').alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('SUPPLY CHAIN MANAGEMENT (SCM) MATRIX BY DASHBOARD ROLE', level=1)
    
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    hdr[0].text = 'Dashboard Role'
    hdr[1].text = 'SCM Functional Area'
    hdr[2].text = 'SCM Business Function'
    hdr[3].text = 'SCM Business Processes'
    for c in hdr:
        c.paragraphs[0].runs[0].font.bold = True

    rows = [
        ('Admin Dashboard\n(/admin/dashboard)', 'Property, Facility & Finance Operations', 'Capacity Planning, Financial Oversight & Machine Asset Control', '• Managing store-wide sales revenue & cash flow.\n• Overseeing machine asset health (idle, washing, maintenance).\n• Managing staff/rider account allocations & system settings.'),
        ('Staff Dashboard\n(/admin/laundry)', 'Operations & Facility Intake', 'Inbound Processing, Machine Dispatch & Quality Assurance', '• Weighing incoming laundry loads (enforcing 7kg max load limit).\n• Scanning QR code tags & controlling stage transitions.\n• Packaging clean laundry & triggering SMS alerts.'),
        ('Rider Dashboard\n(/rider/dashboard)', 'Inbound & Outbound Logistics Distribution', 'Fleet Delivery Management & Doorstep Fulfillment', '• Accepting doorstep pickup requests (out_for_pickup -> received).\n• Transporting laundry between customer & store.\n• Executing delivery (out_for_delivery -> completed).'),
        ('Customer Portal\n(/my-orders & /track)', 'Customer Relationship Management (CRM)', 'Demand Generation, Online Booking & Retention', '• Submitting online laundry bookings (wash, dry, fold, combo).\n• Scanning QR code tags for real-time tracking.\n• Collecting stamps, redeeming ₱50 tokens & rating service.')
    ]

    for r1, r2, r3, r4 in rows:
        row = table.add_row().cells
        row[0].text = r1
        row[1].text = r2
        row[2].text = r3
        row[3].text = r4

    doc.save('Hour_Wash_Subject2_Supply_Chain_Management.docx')
    print('Generated Hour_Wash_Subject2_Supply_Chain_Management.docx')

create_quant_doc()
create_scm_doc()
