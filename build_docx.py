import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_system_design_docx():
    doc = docx.Document()

    # Page Setup - Normal Margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Styles Setup
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A) # Dark slate
    normal_style.paragraph_format.line_spacing = 1.15
    normal_style.paragraph_format.space_after = Pt(6)

    # Document Header / Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_title = p_title.add_run("HOUR WASH LAUNDRY SHOP MANAGEMENT SYSTEM")
    r_title.bold = True
    r_title.font.size = Pt(20)
    r_title.font.color.rgb = RGBColor(0x03, 0x69, 0xA1) # Deep Blue

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_sub = p_sub.add_run("Comprehensive System Architecture, Database Authentication & UML Diagrams Specification (Featuring 38 Dedicated Sequence Diagrams for All UI Navigation Sidebars)")
    r_sub.font.size = Pt(12)
    r_sub.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
    p_sub.paragraph_format.space_after = Pt(16)

    # Horizontal Divider Line
    p_div = doc.add_paragraph()
    r_div = p_div.add_run("―" * 58)
    r_div.font.color.rgb = RGBColor(0xCB, 0xD5, 0xE1)
    p_div.paragraph_format.space_after = Pt(16)

    # ---------------------------------------------------------
    # SECTION 1: SYSTEM DESIGN DIAGRAM
    # ---------------------------------------------------------
    h1 = doc.add_heading(level=1)
    r_h1 = h1.add_run("System Design Diagram")
    r_h1.font.name = 'Arial'
    r_h1.font.size = Pt(16)
    r_h1.bold = True
    r_h1.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    p_disc1 = doc.add_paragraph()
    p_disc1.add_run("System Design & All System Features Breakdown (Isa-Isahin):\n").bold = True
    p_disc1.add_run(
        "A Web-Based Laundry Service Management System for HourWash Laundry Shop in Orosite Legazpi City adopts a multi-tiered, decoupled architecture engineered to deliver real-time operational efficiency, seamless database authentication, and scalable laundry shop management. Built on Laravel 11/12 (PHP 8.5) with Tailwind CSS and Vite, the system structures its capabilities across 13 comprehensive feature modules and 6 architectural layers:\n\n"
        "1. Comprehensive System Feature Breakdown (Isa-Isahin):\n"
        "   • Feature 1: User Authentication & Role-Based Access Control (RBAC): Handles user registration, database login authentication, password hashing (Bcrypt), and session token management for four distinct roles (Customer, Staff, Rider, Admin).\n"
        "   • Feature 2: Customer Profile & Contact Management: Stores customer delivery addresses, barangay, city, province, phone numbers, and specialized laundry instructions in the customer_profiles database table.\n"
        "   • Feature 3: Service Catalog & Pricing Management: Configures laundry service offerings (Wash & Dry, Heavy Duty, Special Detergent, Fabric Conditioner, Fold Only) with weight tariff rates (price per kg/load) and estimated cycle durations in the services table.\n"
        "   • Feature 4: Customer Laundry Order Placement & Scheduling: Allows customers to submit orders, specify load weight estimates, select wash options, choose pickup/delivery types, and set preferred pickup dates and times.\n"
        "   • Feature 5: Washing Machine & Dryer Fleet Management: Tracks real-time machine operational states (Idle, Washing, Rinsing, Drying, Maintenance, Offline), current order assignments, and remaining cycle timers in the machines database table.\n"
        "   • Feature 6: Store Power Outage / Brownout Time Extension (+60 Mins): Enables store staff to trigger a 60-minute time extension on active machine timers during brownouts, adjusting estimated completion timestamps and alerting customers via TextBee SMS.\n"
        "   • Feature 7: Pickup & Delivery Logistics Management: Manages rider assignments, dispatching, real-time status updates (Requested, Scheduled, On the Way, Picked Up, Delivering, Delivered), customer delivery notes, and proof-of-delivery images in the pickup_delivery table.\n"
        "   • Feature 8: Promotional Coupons & Discount Management: Validates promo codes, calculates percentage or flat-rate discounts, updates order subtotal and total_amount fields, and enforces usage limits.\n"
        "   • Feature 9: Unique QR Code Order Verification & Audit Logs (api.qrserver.com): Generates a unique QR token hash and image for every order (qr_codes table) via api.qrserver.com, allowing staff and riders to scan and log verified transactions into the qr_scan_logs audit table.\n"
        "   • Feature 10: Real-Time SMS Phone Notifications (TextBee.dev API): Dispatches automated SMS alerts (SendSmsJob) via TextBee SMS Gateway (api.textbee.dev) upon order status transitions (e.g., Wash Started, Ready for Pickup, Out for Delivery), logging records in sms_notifications.\n"
        "   • Feature 11: Transactional Email Notifications (Brevo API): Sends order confirmation and password reset emails via Brevo Transactional Email Gateway (api.brevo.com).\n"
        "   • Feature 12: Digital Itemized Receipt & Billing: Computes weight-based subtotal, delivery fees, applied discounts, and grand totals, rendering and printing digital transaction receipts.\n"
        "   • Feature 13: Financial Sales Analytics, AI Chatbot (OpenAI / Ollama) & CRM Ratings: Provides administrators with financial revenue charts, 12-stamp Frequent User Card rewards, customer feedback ratings, and integrates an AI Chatbot Assistant powered by OpenAI Cloud LLM (gpt-3.5-turbo) and local Ollama (gemma3:1b).\n\n"
        "2. Multi-Tiered Layer Breakdown:\n"
        "   • Presentation Layer: Blade views, Vite asset bundling, and Tailwind CSS responsive templates matching the exact application sidebars.\n"
        "   • Security & Middleware Layer: AdminMiddleware, StaffMiddleware, CustomerMiddleware, RiderMiddleware, and CSRF protection.\n"
        "   • Application Controller Layer: HTTP request handlers executing core business logic, including ChatbotController.\n"
        "   • Domain & Asynchronous Service Layer: SmsNotificationService (TextBee), EmailNotificationService (Brevo), LoyaltyStampService, and background queue workers.\n"
        "   • Persistence Layer: MySQL relational database with normalized schema tables.\n"
        "   • External Integration Layer: TextBee SMS REST API (api.textbee.dev), Brevo Email REST API (api.brevo.com), OpenAI Cloud API, Ollama LLM, and QRServer Engine (api.qrserver.com)."
    )

    # Insert Figure 1
    p_fig1 = doc.add_paragraph()
    p_fig1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_fig1.paragraph_format.space_before = Pt(12)
    p_fig1.paragraph_format.space_after = Pt(4)
    run_fig1 = p_fig1.add_run()
    run_fig1.add_picture('diagrams/system_design_diagram.png', width=Inches(6.2))

    p_cap1 = doc.add_paragraph()
    p_cap1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_cap1 = p_cap1.add_run("Figure 1: High-Level System Architecture & End-to-End Database Flow Diagram")
    r_cap1.italic = True
    r_cap1.font.size = Pt(9.5)
    r_cap1.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

    p_interp1 = doc.add_paragraph()
    p_interp1.add_run("Figure 1 Interpretation (Step-by-Step Execution Flow):\n").bold = True
    p_interp1.add_run(
        "Figure 1 details how requests flow step-by-step when a user logs in and performs actions. Step 1: User submits login credentials (email and password) via the Presentation Layer. Step 2: The request hits the Routing & Security Layer where CSRF tokens are verified and session middleware intercepts the call. Step 3: AuthenticatedSessionController queries the MySQL database users table to verify the email and check the password hash using Bcrypt. Step 4: Upon database verification, the user's role (Customer, Staff, Rider, Admin) is loaded, a session cookie is generated, and role-specific middleware routes the user to their authorized portal. Step 5: Subsequent actions execute SQL operations on corresponding MySQL tables and enqueue asynchronous notifications via TextBee SMS, Brevo Email, QRServer API, and OpenAI/Ollama LLM."
    )

    # ---------------------------------------------------------
    # SECTION 2: USE CASE DIAGRAM
    # ---------------------------------------------------------
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    h2_uc = doc.add_heading(level=1)
    r_h2_uc = h2_uc.add_run("Use Case Diagram")
    r_h2_uc.font.name = 'Arial'
    r_h2_uc.font.size = Pt(16)
    r_h2_uc.bold = True
    r_h2_uc.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    p_disc2 = doc.add_paragraph()
    p_disc2.add_run("Use Case Diagram Discussion & Actor Responsibilities (Matching All 38 UI Navigation Menu Items):\n").bold = True
    p_disc2.add_run(
        "The Use Case Diagram defines the functional boundary of the HourWash system across 38 distinct use cases matching all navigation features from the 4 application UI screenshots in exact order:\n\n"
        "• Customer Actor Navigation Use Cases (Screenshot 3 Sidebar):\n"
        "   - UC1: Customer Registration\n"
        "   - UC2: Customer Login Authentication\n"
        "   - UC3: Customer Forgot Password Reset (Brevo API)\n"
        "   - UC4: Customer Dashboard (Live status & quick actions)\n"
        "   - UC5: Book New Order (New wash booking request)\n"
        "   - UC6: My Order History (Track bookings & receipts via api.qrserver.com)\n"
        "   - UC7: Frequent User Card (12-stamp loyalty rewards)\n"
        "   - UC8: Home Dashboard (Public landing page)\n"
        "   - UC9: Account Settings (Profile & security settings)\n\n"
        "• Staff Operator Actor Navigation Use Cases (Screenshot 2 Sidebar):\n"
        "   - UC10: Staff Login Authentication\n"
        "   - UC11: Staff Forgot Password Reset\n"
        "   - UC12: Workstation Dashboard (Queue & cashier processing overview)\n"
        "   - UC13: Manage Laundry Orders (Queue & cashier processing)\n"
        "   - UC14: Manage Machines (Add, Edit, & Remove Machines / +60m Extension)\n"
        "   - UC15: New Walk-in Order (Book customer wash at store counter)\n"
        "   - UC16: QR Scan Logs Outbox (Audit log of all QR scans)\n"
        "   - UC17: Home Dashboard (Public landing page)\n"
        "   - UC18: Account Settings (Profile & security settings)\n\n"
        "• Rider Logistics Actor Navigation Use Cases (Screenshot 4 Sidebar):\n"
        "   - UC19: Rider Login Authentication\n"
        "   - UC20: Rider Forgot Password Reset\n"
        "   - UC21: Rider of Hour Wash (Pickup & delivery dispatches)\n"
        "   - UC22: Update Pickup Logistics Status ('On the Way', 'Picked Up' via TextBee SMS)\n"
        "   - UC23: Update Delivery Status & Upload Proof Photo Image ('Delivered')\n"
        "   - UC24: Home Dashboard (Public landing page)\n"
        "   - UC25: Account Settings (Profile & security settings)\n\n"
        "• System Administrator Actor Navigation Use Cases (Screenshot 1 Sidebar):\n"
        "   - UC26: Admin Login Authentication\n"
        "   - UC27: Admin Forgot Password Reset\n"
        "   - UC28: Overall Reports & Dashboard (System overview & metrics)\n"
        "   - UC29: Manage Laundry Orders (Queue & cashier processing)\n"
        "   - UC30: Manage Machines (Add, Edit, & Remove Machines)\n"
        "   - UC31: Services & Pricing (Service rates & load options)\n"
        "   - UC32: Manage Users (Stamps, Add, Edit, & Remove Users)\n"
        "   - UC33: Live SMS Outbox (Phone notification logs via TextBee Gateway)\n"
        "   - UC34: Live Email Outbox (Email notification logs via Brevo Gateway)\n"
        "   - UC35: Customer Reviews Outbox (Ratings & feedback logs)\n"
        "   - UC36: QR Scan Logs Outbox (Audit log of all QR scans)\n"
        "   - UC37: Home Dashboard (Public landing page)\n"
        "   - UC38: Account Settings (Profile & security settings)"
    )

    p_fig2 = doc.add_paragraph()
    p_fig2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_fig2.paragraph_format.space_before = Pt(12)
    p_fig2.paragraph_format.space_after = Pt(4)
    run_fig2 = p_fig2.add_run()
    run_fig2.add_picture('diagrams/use_case_diagram.png', width=Inches(6.2))

    p_cap2 = doc.add_paragraph()
    p_cap2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_cap2 = p_cap2.add_run("Figure 2: HourWash System Use Case Diagram - Matching All 38 UI Navigation Sidebar Items")
    r_cap2.italic = True
    r_cap2.font.size = Pt(9.5)
    r_cap2.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

    p_interp2 = doc.add_paragraph()
    p_interp2.add_run("Figure 2 Interpretation:\n").bold = True
    p_interp2.add_run(
        "Figure 2 visually maps all 38 use cases enclosed within the system boundary box. Associations connect each actor to their permissible interactions matching the exact navigation bar menus from the application UI."
    )

    # ---------------------------------------------------------
    # SECTION 3: CLASS DIAGRAM
    # ---------------------------------------------------------
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    h3_cd = doc.add_heading(level=1)
    r_h3_cd = h3_cd.add_run("Class Diagram")
    r_h3_cd.font.name = 'Arial'
    r_h3_cd.font.size = Pt(16)
    r_h3_cd.bold = True
    r_h3_cd.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    p_disc3 = doc.add_paragraph()
    p_disc3.add_run("Class Diagram Discussion & Relational Database Mapping (Isa-Isahin):\n").bold = True
    p_disc3.add_run(
        "The Class Diagram presents the static object-oriented domain model of the HourWash application, mapping directly to Laravel Eloquent ORM models and MySQL database tables:\n\n"
        "• User Model (users table): Attributes: id (PK), name, email, password, role ('customer', 'staff', 'rider', 'admin'), frequent_stamps. Methods: authenticate(), claimLoyaltyStamp().\n"
        "• CustomerProfile Model (customer_profiles table): Attributes: id (PK), user_id (FK), address, barangay, city, province. BelongsTo User.\n"
        "• StaffProfile Model (staff_profiles table): Attributes: id (PK), user_id (FK), employee_id, position, status. BelongsTo User.\n"
        "• Service Model (services table): Attributes: id (PK), name, service_type, price, estimated_minutes. HasMany Order.\n"
        "• Machine Model (machines table): Attributes: id (PK), machine_code, machine_type, status, current_order_id (FK). Methods: assignOrder(), addBrownoutTime().\n"
        "• Order Model (orders table): Attributes: id (PK), order_number, customer_id (FK), service_id (FK), total_amount, order_status. HasMany OrderStatusHistory, HasOne QrCode, HasOne PickupDelivery.\n"
        "• OrderStatusHistory Model (order_status_history table): Attributes: id (PK), order_id (FK), status, changed_by (FK), created_at.\n"
        "• PickupDelivery Model (pickup_delivery table): Attributes: id (PK), order_id (FK), type, status, proof_images. BelongsTo Order.\n"
        "• QrCode & QrScanLog Models (qr_codes & qr_scan_logs tables): QrCode stores order_id (FK), qr_token. QrScanLog stores scanned_by (FK), scan_type.\n"
        "• SmsNotification & EmailNotification Models (sms_notifications & email_notifications tables): Stores dispatch logs for TextBee SMS and Brevo Email APIs."
    )

    p_fig3 = doc.add_paragraph()
    p_fig3.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_fig3.paragraph_format.space_before = Pt(12)
    p_fig3.paragraph_format.space_after = Pt(4)
    run_fig3 = p_fig3.add_run()
    run_fig3.add_picture('diagrams/class_diagram.png', width=Inches(6.2))

    p_cap3 = doc.add_paragraph()
    p_cap3.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_cap3 = p_cap3.add_run("Figure 3: HourWash System Complete Domain Class Diagram & Relational Schema")
    r_cap3.italic = True
    r_cap3.font.size = Pt(9.5)
    r_cap3.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

    p_interp3 = doc.add_paragraph()
    p_interp3.add_run("Figure 3 Interpretation:\n").bold = True
    p_interp3.add_run(
        "Figure 3 depicts entity structures, attribute data types, primary/foreign key mappings, and association multiplicities within the HourWash database domain."
    )

    # ---------------------------------------------------------
    # SECTION 4: SEQUENCE DIAGRAMS (UC1 to UC38)
    # ---------------------------------------------------------
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    h4_sd = doc.add_heading(level=1)
    r_h4_sd = h4_sd.add_run("Sequence Diagram")
    r_h4_sd.font.name = 'Arial'
    r_h4_sd.font.size = Pt(16)
    r_h4_sd.bold = True
    r_h4_sd.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    p_disc4 = doc.add_paragraph()
    p_disc4.add_run("Sequence Diagram Discussion:\n").bold = True
    p_disc4.add_run(
        "Sequence diagrams model the dynamic execution flows and step-by-step ('isa-isahin') message passing between user lifelines, controllers, models, and MySQL database tables during critical system operations for all 38 Use Cases across the four navigation sidebars."
    )

    # 38 Sequence Diagram items definitions matching UI navigation
    SD_ITEMS_38 = [
        # Customer Role (UC1 - UC9)
        ("Customer Role Navigation Workflows (Screenshot 3 Sidebar)", 1, "UC1: Customer Account Registration Flow", "sequence_diagram_1.png",
         "Figure 4: Sequence Diagram 1 - UC1 Customer Account Registration Flow",
         "Figure 4 details customer registration. The user submits registration details, RegisteredUserController validates input, creates a row in the users table (role='customer'), and inserts contact metadata into customer_profiles."),

        (None, 2, "UC2: Customer Login Authentication Flow", "sequence_diagram_2.png",
         "Figure 5: Sequence Diagram 2 - UC2 Customer Login Authentication Flow",
         "Figure 5 details customer authentication. The customer inputs email and password, AuthenticatedSessionController queries the users table, Hash::check() validates the password hash, and CustomerMiddleware grants portal access."),

        (None, 3, "UC3: Customer Forgot Password Reset Flow via Brevo Email API", "sequence_diagram_3.png",
         "Figure 6: Sequence Diagram 3 - UC3 Customer Forgot Password Reset Flow via Brevo Email API",
         "Figure 6 details customer password recovery. PasswordResetLinkController passes credentials to EmailNotificationService, which executes an HTTP POST call to Brevo API (https://api.brevo.com/v3/smtp/email)."),

        (None, 4, "UC4: Customer Dashboard Navigation Flow", "sequence_diagram_4.png",
         "Figure 7: Sequence Diagram 4 - UC4 Customer Dashboard Navigation Flow",
         "Figure 7 details Customer Dashboard navigation. CustomerDashboardController queries active bookings and machine states to render live order status timelines and quick action buttons."),

        (None, 5, "UC5: Book New Order Navigation Flow", "sequence_diagram_5.png",
         "Figure 8: Sequence Diagram 5 - UC5 Book New Order Navigation Flow",
         "Figure 8 models new wash bookings. The customer selects service options and load weight. LaundryController queries Service model rates and inserts an order record into the MySQL orders table."),

        (None, 6, "UC6: My Order History Navigation Flow & QR Rendering", "sequence_diagram_6.png",
         "Figure 9: Sequence Diagram 6 - UC6 My Order History Navigation Flow & QR Rendering",
         "Figure 9 details order history tracking. OrderHistoryController fetches past bookings and qr_token, rendering a digital QR code image dynamically via api.qrserver.com (https://api.qrserver.com/v1/create-qr-code/)."),

        (None, 7, "UC7: Frequent User Card (12-Stamp Loyalty Rewards) Navigation Flow", "sequence_diagram_7.png",
         "Figure 10: Sequence Diagram 7 - UC7 Frequent User Card Navigation Flow",
         "Figure 10 models 12-stamp loyalty card tracking. LoyaltyCardController queries user frequent_user_card stamp count from MySQL users table, evaluating reward eligibility for free wash services."),

        (None, 8, "UC8: Home Dashboard Public Landing Page Navigation Flow", "sequence_diagram_8.png",
         "Figure 11: Sequence Diagram 8 - UC8 Home Dashboard Navigation Flow",
         "Figure 11 details home landing page rendering. HomeController queries public active services, pricing rates, and store operating hours to render the landing page view."),

        (None, 9, "UC9: Account Settings Navigation Flow", "sequence_diagram_9.png",
         "Figure 12: Sequence Diagram 9 - UC9 Account Settings Navigation Flow",
         "Figure 12 depicts profile updates. The customer modifies contact info or password, and ProfileController executes SQL UPDATE statements on users and customer_profiles tables."),

        # Staff Role (UC10 - UC18)
        ("Staff Operator Role Navigation Workflows (Screenshot 2 Sidebar)", 10, "UC10: Staff Login Authentication Flow", "sequence_diagram_10.png",
         "Figure 13: Sequence Diagram 10 - UC10 Staff Login Authentication Flow",
         "Figure 13 details staff authentication. The staff operator inputs credentials, AuthenticatedSessionController verifies the user record (role='staff'), validates the password hash, and StaffMiddleware grants console access."),

        (None, 11, "UC11: Staff Forgot Password Reset Flow via Brevo Email API", "sequence_diagram_11.png",
         "Figure 14: Sequence Diagram 11 - UC11 Staff Forgot Password Reset Flow via Brevo Email API",
         "Figure 14 details staff password recovery. PasswordResetLinkController invokes EmailNotificationService to dispatch a secure reset link via Brevo Email API (api.brevo.com)."),

        (None, 12, "UC12: Workstation Dashboard Navigation Flow", "sequence_diagram_12.png",
         "Figure 15: Sequence Diagram 12 - UC12 Workstation Dashboard Navigation Flow",
         "Figure 15 models Workstation Dashboard navigation. WorkstationController fetches active laundry queue items, machine timer statuses, and cashier counter metrics."),

        (None, 13, "UC13: Manage Laundry Orders Navigation Flow", "sequence_diagram_13.png",
         "Figure 16: Sequence Diagram 13 - UC13 Manage Laundry Orders Navigation Flow",
         "Figure 16 models load weighing. Staff inputs physical laundry scale weight, LaundryController calculates subtotal = weight * rate, and updates order status to 'weighed' in MySQL."),

        (None, 14, "UC14: Manage Machines Navigation Flow (+60m Extension & TextBee Alert)", "sequence_diagram_14.png",
         "Figure 17: Sequence Diagram 14 - UC14 Manage Machines Navigation Flow",
         "Figure 17 details machine management. Staff triggers a +60m brownout extension on active machines, updating remaining_minutes in DB and dispatching an automated SMS alert via TextBee SMS Gateway (api.textbee.dev)."),

        (None, 15, "UC15: New Walk-in Order Navigation Flow", "sequence_diagram_15.png",
         "Figure 18: Sequence Diagram 15 - UC15 New Walk-in Order Navigation Flow",
         "Figure 18 details walk-in order creation. Staff enters customer details at the store counter, creating or finding the user record and creating an order in the database."),

        (None, 16, "UC16: QR Scan Logs Outbox Navigation Flow", "sequence_diagram_16.png",
         "Figure 19: Sequence Diagram 16 - UC16 QR Scan Logs Outbox Navigation Flow",
         "Figure 19 models QR scan log inspection. QrScanLogController queries qr_scan_logs audit table to render scan event history for staff inspection."),

        (None, 17, "UC17: Staff Home Dashboard Navigation Flow", "sequence_diagram_17.png",
         "Figure 20: Sequence Diagram 17 - UC17 Staff Home Dashboard Navigation Flow",
         "Figure 20 details staff home page access. HomeController queries active service metadata to render the public store landing view."),

        (None, 18, "UC18: Staff Account Settings Navigation Flow", "sequence_diagram_18.png",
         "Figure 21: Sequence Diagram 18 - UC18 Staff Account Settings Navigation Flow",
         "Figure 21 details staff profile updates. The staff member updates password or security settings, executing SQL UPDATE statements on the users table."),

        # Rider Role (UC19 - UC25)
        ("Rider Logistics Role Navigation Workflows (Screenshot 4 Sidebar)", 19, "UC19: Rider Login Authentication Flow", "sequence_diagram_19.png",
         "Figure 22: Sequence Diagram 19 - UC19 Rider Login Authentication Flow",
         "Figure 22 details rider authentication. The rider inputs login credentials, AuthenticatedSessionController verifies the user record (role='rider'), checks the password hash, and RiderMiddleware grants access."),

        (None, 20, "UC20: Rider Forgot Password Reset Flow via Brevo Email API", "sequence_diagram_20.png",
         "Figure 23: Sequence Diagram 20 - UC20 Rider Forgot Password Reset Flow via Brevo Email API",
         "Figure 23 details rider password recovery. PasswordResetLinkController triggers EmailNotificationService to dispatch a reset email via Brevo Email Gateway."),

        (None, 21, "UC21: Rider of Hour Wash Navigation Flow", "sequence_diagram_21.png",
         "Figure 24: Sequence Diagram 21 - UC21 Rider of Hour Wash Navigation Flow",
         "Figure 24 models rider task access. The rider opens the Rider Dashboard, and RiderDashboardController queries the pickup_delivery table for assigned jobs."),

        (None, 22, "UC22: Update Pickup Logistics Status Flow & TextBee SMS Alert", "sequence_diagram_22.png",
         "Figure 25: Sequence Diagram 22 - UC22 Update Pickup Logistics Status Flow",
         "Figure 25 details pickup completion. Rider updates status to 'picked_up', and SmsNotificationService dispatches an automated SMS alert to customer phone via TextBee Gateway (api.textbee.dev)."),

        (None, 23, "UC23: Update Delivery Status & Proof Photo Upload Flow", "sequence_diagram_23.png",
         "Figure 26: Sequence Diagram 23 - UC23 Update Delivery Status & Proof Photo Upload Flow",
         "Figure 26 models delivery completion. The rider uploads a proof photo, PickupDeliveryController updates status to 'delivered', sets order to 'completed', and alerts the customer."),

        (None, 24, "UC24: Rider Home Dashboard Navigation Flow", "sequence_diagram_24.png",
         "Figure 27: Sequence Diagram 24 - UC24 Rider Home Dashboard Navigation Flow",
         "Figure 27 details rider home landing page access, rendering store hours and active service packages."),

        (None, 25, "UC25: Rider Account Settings Navigation Flow", "sequence_diagram_25.png",
         "Figure 28: Sequence Diagram 25 - UC25 Rider Account Settings Navigation Flow",
         "Figure 28 details rider profile updates. The rider updates profile details or password, executing SQL UPDATE statements on the users table."),

        # Admin Role (UC26 - UC38)
        ("Administrator Role Navigation Workflows (Screenshot 1 Sidebar)", 26, "UC26: Administrator Login Authentication Flow", "sequence_diagram_26.png",
         "Figure 29: Sequence Diagram 26 - UC26 Administrator Login Authentication Flow",
         "Figure 29 details admin authentication. The administrator inputs credentials, AuthenticatedSessionController verifies the user record (role='admin'), checks password hash, and AdminMiddleware grants access."),

        (None, 27, "UC27: Administrator Password Reset Flow via Brevo Email API", "sequence_diagram_27.png",
         "Figure 30: Sequence Diagram 27 - UC27 Administrator Password Reset Flow via Brevo Email API",
         "Figure 30 details admin password recovery. PasswordResetLinkController triggers EmailNotificationService to send a secure reset link via Brevo Email API."),

        (None, 28, "UC28: Overall Reports & Dashboard Navigation Flow", "sequence_diagram_28.png",
         "Figure 31: Sequence Diagram 28 - UC28 Overall Reports & Dashboard Navigation Flow",
         "Figure 31 details Overall Reports Dashboard rendering. AnalyticsController aggregates overall sales revenue, daily profit, and machine utilization metrics."),

        (None, 29, "UC29: Manage Laundry Orders Navigation Flow", "sequence_diagram_29.png",
         "Figure 32: Sequence Diagram 29 - UC29 Manage Laundry Orders Navigation Flow",
         "Figure 32 shows admin order management. Admin reviews order queue status and overrides order statuses, updating orders and order_status_history tables."),

        (None, 30, "UC30: Manage Machines Navigation Flow (Add, Edit, Remove)", "sequence_diagram_30.png",
         "Figure 33: Sequence Diagram 30 - UC30 Manage Machines Navigation Flow",
         "Figure 33 models machine fleet management. Admin adds, edits, or removes machines, executing SQL INSERT, UPDATE, or DELETE statements on the machines table."),

        (None, 31, "UC31: Services & Pricing Navigation Flow", "sequence_diagram_31.png",
         "Figure 34: Sequence Diagram 31 - UC31 Services & Pricing Navigation Flow",
         "Figure 34 details service rate management. Admin configures pricing tariffs per kg and estimated cycle minutes, updating the services table."),

        (None, 32, "UC32: Manage Users Navigation Flow (Stamps, Add, Edit, Remove)", "sequence_diagram_32.png",
         "Figure 35: Sequence Diagram 32 - UC32 Manage Users Navigation Flow",
         "Figure 35 details user & stamp management. Admin manages user profiles and adjusts 12-stamp Frequent User Card counts in the users table."),

        (None, 33, "UC33: Live SMS Outbox Navigation Flow (TextBee Logs)", "sequence_diagram_33.png",
         "Figure 36: Sequence Diagram 33 - UC33 Live SMS Outbox Navigation Flow",
         "Figure 36 details Live TextBee SMS Outbox inspection. SmsLogController queries sms_notifications table to display phone SMS delivery logs from TextBee Gateway."),

        (None, 34, "UC34: Live Email Outbox Navigation Flow (Brevo Logs)", "sequence_diagram_34.png",
         "Figure 37: Sequence Diagram 34 - UC34 Live Brevo Email Outbox Navigation Flow",
         "Figure 37 details Live Brevo Email Outbox inspection. EmailLogController queries email_notifications table to render transactional email logs from Brevo Gateway."),

        (None, 35, "UC35: Customer Reviews Outbox Navigation Flow", "sequence_diagram_35.png",
         "Figure 38: Sequence Diagram 35 - UC35 Customer Reviews Outbox Navigation Flow",
         "Figure 38 shows customer review management. CustomerFeedbackController queries ratings and comments from customer_feedbacks table."),

        (None, 36, "UC36: QR Scan Logs Outbox Navigation Flow", "sequence_diagram_36.png",
         "Figure 39: Sequence Diagram 36 - UC36 QR Scan Logs Outbox Navigation Flow",
         "Figure 39 models global QR audit log inspection. QrScanLogController queries all scan events from qr_scan_logs table for administrative auditing."),

        (None, 37, "UC37: Admin Home Dashboard Navigation Flow", "sequence_diagram_37.png",
         "Figure 40: Sequence Diagram 37 - UC37 Admin Home Dashboard Navigation Flow",
         "Figure 40 details admin home page navigation, rendering public landing page metrics and store info."),

        (None, 38, "UC38: Admin Account Settings Navigation Flow", "sequence_diagram_38.png",
         "Figure 41: Sequence Diagram 38 - UC38 Admin Account Settings Navigation Flow",
         "Figure 41 details admin profile updates. Admin updates password or security settings, executing SQL UPDATE statements on the users table."),
    ]

    for section_header, sd_num, uc_title, img_filename, cap_text, interp_text in SD_ITEMS_38:
        if section_header:
            doc.add_paragraph().paragraph_format.space_after = Pt(10)
            h_sec = doc.add_heading(level=2)
            r_h_sec = h_sec.add_run(f"=== {section_header.upper()} ===")
            r_h_sec.font.name = 'Arial'
            r_h_sec.font.size = Pt(14)
            r_h_sec.bold = True
            r_h_sec.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

        doc.add_paragraph().paragraph_format.space_after = Pt(4)
        h_sd = doc.add_heading(level=3)
        r_h_sd = h_sd.add_run(f"Sequence Diagram {sd_num}")
        r_h_sd.font.name = 'Arial'
        r_h_sd.font.size = Pt(12.5)
        r_h_sd.bold = True
        r_h_sd.font.color.rgb = RGBColor(0x03, 0x69, 0xA1)

        p_align = doc.add_paragraph()
        p_align.add_run("Alignment with Use Case Diagram Feature/Process Name:\n").bold = True
        p_align.add_run(uc_title)

        p_fig = doc.add_paragraph()
        p_fig.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_fig.paragraph_format.space_before = Pt(8)
        p_fig.paragraph_format.space_after = Pt(4)
        run_fig = p_fig.add_run()
        run_fig.add_picture(f'diagrams/{img_filename}', width=Inches(6.2))

        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r_cap = p_cap.add_run(cap_text)
        r_cap.italic = True
        r_cap.font.size = Pt(9.5)
        r_cap.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

        p_interp = doc.add_paragraph()
        p_interp.add_run(f"Figure {sd_num + 3} Interpretation:\n").bold = True
        p_interp.add_run(interp_text)

    # ---------------------------------------------------------
    # SECTION 5: PACKAGE DIAGRAM
    # ---------------------------------------------------------
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    h5_pd = doc.add_heading(level=1)
    r_h5_pd = h5_pd.add_run("Package Diagram")
    r_h5_pd.font.name = 'Arial'
    r_h5_pd.font.size = Pt(16)
    r_h5_pd.bold = True
    r_h5_pd.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    p_disc5 = doc.add_paragraph()
    p_disc5.add_run("Package Diagram Discussion & Subsystem Architecture (Isa-Isahin):\n").bold = True
    p_disc5.add_run(
        "The Package Diagram structures the HourWash codebase into modular namespaces under App\\ and Database\\:\n\n"
        "• App\\Http\\Controllers: Contains Auth\\AuthenticatedSessionController, LaundryController, MachineController, ChatbotController (OpenAI / Ollama AI), RiderDashboardController, and AnalyticsController.\n"
        "• App\\Http\\Middleware: Contains AdminMiddleware, StaffMiddleware, CustomerMiddleware, RiderMiddleware, and SecurityHeaders.\n"
        "• App\\Models: Contains User (frequent_user_card), CustomerProfile, StaffProfile, Order, Machine, PickupDelivery, QrCode, QrScanLog, SmsNotification, and CustomerFeedback.\n"
        "• App\\Services: Contains SmsNotificationService (TextBee.dev API) and EmailNotificationService (Brevo API).\n"
        "• App\\Jobs & App\\Mail: Contains SendSmsJob (TextBee) and OrderStatusUpdated Mail (Brevo).\n"
        "• Database\\Migrations: Contains migration files creating MySQL database tables.\n"
        "• Resources\\Views: Contains server-side Blade templates organized by role (auth/*, customer/*, staff/*, rider/*, admin/*)."
    )

    p_fig42 = doc.add_paragraph()
    p_fig42.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_fig42.paragraph_format.space_before = Pt(12)
    p_fig42.paragraph_format.space_after = Pt(4)
    run_fig42 = p_fig42.add_run()
    run_fig42.add_picture('diagrams/package_diagram.png', width=Inches(6.2))

    p_cap42 = doc.add_paragraph()
    p_cap42.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_cap42 = p_cap42.add_run("Figure 42: HourWash System Subsystem Package Diagram")
    r_cap42.italic = True
    r_cap42.font.size = Pt(9.5)
    r_cap42.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

    p_interp42 = doc.add_paragraph()
    p_interp42.add_run("Figure 42 Interpretation:\n").bold = True
    p_interp42.add_run(
        "Figure 42 details the clean packaging layout of the application. Dependencies flow inwards from HTTP Controllers through Security Middleware down to Eloquent Models and Service classes, preserving modularity and low coupling."
    )

    # ---------------------------------------------------------
    # SECTION 6: DEPLOYMENT DIAGRAM
    # ---------------------------------------------------------
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    h6_dd = doc.add_heading(level=1)
    r_h6_dd = h6_dd.add_run("Deployment Diagram")
    r_h6_dd.font.name = 'Arial'
    r_h6_dd.font.size = Pt(16)
    r_h6_dd.bold = True
    r_h6_dd.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    p_disc6 = doc.add_paragraph()
    p_disc6.add_run("Deployment Diagram Discussion & Node Topology (Isa-Isahin):\n").bold = True
    p_disc6.add_run(
        "The Deployment Diagram details the hardware nodes, execution runtimes, and network protocols supporting the HourWash system in production:\n\n"
        "1. Client Devices Node: Mobile smartphones and desktop web browsers running Tailwind CSS user interfaces and executing Vite/JavaScript code.\n"
        "2. Web Application Server Node: Ubuntu Linux or Windows Server hosting Nginx/Apache Web Server, PHP 8.5 FPM runtime, Laravel 11 application engine, and Artisan task queue worker.\n"
        "3. Database Server Node: Dedicated database node running MySQL Server 8.0 with InnoDB storage engine, storing encrypted user password hashes and relational data tables.\n"
        "4. External Integration Gateways: TextBee SMS REST API (api.textbee.dev), Brevo Email REST API (api.brevo.com), OpenAI Cloud API (api.openai.com), local Ollama LLM (127.0.0.1:11434), and QRServer Engine (api.qrserver.com)."
    )

    p_fig43 = doc.add_paragraph()
    p_fig43.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_fig43.paragraph_format.space_before = Pt(12)
    p_fig43.paragraph_format.space_after = Pt(4)
    run_fig43 = p_fig43.add_run()
    run_fig43.add_picture('diagrams/deployment_diagram.png', width=Inches(6.2))

    p_cap43 = doc.add_paragraph()
    p_cap43.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_cap43 = p_cap43.add_run("Figure 43: HourWash Production Infrastructure & Deployment Diagram")
    r_cap43.italic = True
    r_cap43.font.size = Pt(9.5)
    r_cap43.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

    p_interp43 = doc.add_paragraph()
    p_interp43.add_run("Figure 43 Interpretation:\n").bold = True
    p_interp43.add_run(
        "Figure 43 illustrates the network communication channels connecting hardware nodes. Client Devices communicate with the Web Application Server via secure HTTPS (Port 443). The Application Server executes PDO SQL queries against the Database Server Node over TCP/IP Port 3306. Outbound REST API calls (HTTPS Port 443) route notification and AI requests to TextBee SMS, Brevo Email, OpenAI Cloud, and QRServer APIs."
    )

    output_filename = "Hour_Wash_System_Design_Diagrams.docx"
    doc.save(output_filename)
    print(f"DOCUMENT SUCCESSFULLY CREATED WITH ALL 38 SEQUENCE DIAGRAMS AND SAVED TO: {output_filename}")

if __name__ == '__main__':
    create_system_design_docx()
